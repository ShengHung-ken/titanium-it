from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent

INPUT_FILE = (
    BASE_DIR
    / "鈦鼎資訊_Titanium_IT_官網建置與維運SOP_V1.0.docx"
)

OUTPUT_FILE = (
    BASE_DIR
    / "鈦鼎資訊_Titanium_IT_官網建置與維運SOP_V1.1.docx"
)


def set_run_font(
    run,
    chinese_font="Microsoft JhengHei",
    english_font="Aptos",
    size=11,
    bold=False,
    color=None,
):
    run.font.name = english_font
    run.font.size = Pt(size)
    run.bold = bold

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        chinese_font,
    )

    if color:
        run.font.color.rgb = RGBColor(
            *color
        )


def format_paragraph(
    paragraph,
    before=0,
    after=6,
    line_spacing=1.25,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def insert_heading_before(
    anchor,
    text,
    level=1,
):
    style_name = f"Heading {level}"

    paragraph = anchor.insert_paragraph_before(
        text,
        style=style_name,
    )

    format_paragraph(
        paragraph,
        before=8,
        after=8,
    )

    return paragraph


def insert_text_before(
    anchor,
    text,
    bold=False,
):
    paragraph = anchor.insert_paragraph_before()

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=11,
        bold=bold,
    )

    format_paragraph(
        paragraph,
        after=6,
    )

    return paragraph


def insert_bullet_before(
    anchor,
    text,
):
    paragraph = anchor.insert_paragraph_before()

    run = paragraph.add_run(
        f"• {text}"
    )

    set_run_font(
        run,
        size=11,
    )

    paragraph.paragraph_format.left_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.first_line_indent = Cm(
        -0.35
    )

    format_paragraph(
        paragraph,
        after=3,
    )

    return paragraph


def insert_numbered_before(
    anchor,
    number,
    text,
):
    paragraph = anchor.insert_paragraph_before()

    run = paragraph.add_run(
        f"{number}. {text}"
    )

    set_run_font(
        run,
        size=11,
    )

    paragraph.paragraph_format.left_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.first_line_indent = Cm(
        -0.45
    )

    format_paragraph(
        paragraph,
        after=4,
    )

    return paragraph


def insert_code_before(
    anchor,
    text,
):
    paragraph = anchor.insert_paragraph_before()

    paragraph.paragraph_format.left_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.right_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.space_before = Pt(
        4
    )
    paragraph.paragraph_format.space_after = Pt(
        8
    )

    p_pr = paragraph._p.get_or_add_pPr()

    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        "F3F5F7",
    )
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")

    for side in (
        "top",
        "left",
        "bottom",
        "right",
    ):
        border = OxmlElement(
            f"w:{side}"
        )

        border.set(
            qn("w:val"),
            "single",
        )
        border.set(
            qn("w:sz"),
            "4",
        )
        border.set(
            qn("w:color"),
            "D9E1E8",
        )

        borders.append(border)

    p_pr.append(borders)

    run = paragraph.add_run(text)

    set_run_font(
        run,
        chinese_font="Microsoft JhengHei",
        english_font="Consolas",
        size=9.5,
        color=(30, 40, 50),
    )

    return paragraph


def insert_note_before(
    anchor,
    title,
    text,
    level="info",
):
    colors = {
        "info": (
            "EAF3FF",
            (47, 111, 179),
        ),
        "warning": (
            "FFF4CC",
            (183, 121, 31),
        ),
        "danger": (
            "FFE5E5",
            (180, 35, 24),
        ),
    }

    fill_color, title_color = colors.get(
        level,
        colors["info"],
    )

    paragraph = anchor.insert_paragraph_before()

    p_pr = paragraph._p.get_or_add_pPr()

    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        fill_color,
    )

    p_pr.append(shading)

    paragraph.paragraph_format.left_indent = Cm(
        0.3
    )
    paragraph.paragraph_format.right_indent = Cm(
        0.3
    )
    paragraph.paragraph_format.space_before = Pt(
        5
    )
    paragraph.paragraph_format.space_after = Pt(
        8
    )

    title_run = paragraph.add_run(
        f"{title}\n"
    )

    set_run_font(
        title_run,
        size=11,
        bold=True,
        color=title_color,
    )

    body_run = paragraph.add_run(text)

    set_run_font(
        body_run,
        size=10.5,
    )

    return paragraph


def find_paragraph_exact(
    document,
    text,
):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph

    return None


def update_cover_version(
    document,
):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text == "文件版本：V1.0":
            for run in paragraph.runs:
                if "V1.0" in run.text:
                    run.text = run.text.replace(
                        "V1.0",
                        "V1.1",
                    )

        if text == "官網建置與維運 SOP V1.0":
            for run in paragraph.runs:
                if "V1.0" in run.text:
                    run.text = run.text.replace(
                        "V1.0",
                        "V1.1",
                    )


def update_document_info_table(
    document,
):
    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue

            key = row.cells[0].text.strip()

            if key == "文件版本":
                row.cells[1].text = "V1.1"

                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            size=10,
                        )

                return


def add_version_history(
    document,
):
    for table in document.tables:
        if not table.rows:
            continue

        header = [
            cell.text.strip()
            for cell in table.rows[0].cells
        ]

        if header == [
            "版本",
            "日期",
            "修改內容",
        ]:
            # 避免重複新增
            for row in table.rows[1:]:
                if (
                    row.cells[0].text.strip()
                    == "V1.1"
                ):
                    return

            cells = table.add_row().cells

            cells[0].text = "V1.1"
            cells[1].text = "2026-08-28"
            cells[2].text = (
                "新增「新電腦移機 / "
                "開發環境還原 SOP」，"
                "補充舊電腦移機前檢查、"
                "新電腦環境重建、GitHub Clone、"
                ".env.local 還原、npm 安裝、"
                "Build 驗證及雲端服務還原說明。"
            )

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            size=10,
                        )

            return


def enable_field_update(
    document,
):
    settings = document.settings.element

    for element in settings.findall(
        qn("w:updateFields")
    ):
        settings.remove(element)

    update_fields = OxmlElement(
        "w:updateFields"
    )

    update_fields.set(
        qn("w:val"),
        "true",
    )

    settings.append(update_fields)


# ============================================================
# 檢查來源檔
# ============================================================

if not INPUT_FILE.exists():
    print()
    print("找不到原始 SOP：")
    print(INPUT_FILE)
    print()
    print(
        "請確認 V1.0 Word 檔與本程式"
        "放在同一個資料夾。"
    )
    raise SystemExit(1)


# ============================================================
# 開啟 V1.0
# ============================================================

document = Document(
    str(INPUT_FILE)
)


# ============================================================
# 找到原本「第 26 章 版本修改紀錄」
# ============================================================

version_heading = find_paragraph_exact(
    document,
    "第 26 章　版本修改紀錄",
)

if version_heading is None:
    print()
    print(
        "找不到「第 26 章　版本修改紀錄」。"
    )
    print(
        "請確認使用的是原本產生的 V1.0 SOP。"
    )
    raise SystemExit(1)


# ============================================================
# 插入換頁
# ============================================================

page_break = (
    version_heading.insert_paragraph_before()
)

page_break.add_run().add_break(
    WD_BREAK.PAGE
)


# ============================================================
# 第 26 章
# ============================================================

insert_heading_before(
    version_heading,
    "第 26 章　新電腦移機 / 開發環境還原 SOP",
    level=1,
)


# ------------------------------------------------------------
# 26.1
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.1 目的",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "本章用於舊電腦更換、新電腦接手、"
        "硬碟故障或重新安裝 Windows 後，"
        "快速還原 Titanium IT 官方網站的"
        "完整開發環境。"
    ),
)

insert_note_before(
    version_heading,
    "核心觀念",
    (
        "GitHub 保存主程式；Supabase 保存商品、"
        "帳號與圖片資料；新電腦主要只需要重新安裝"
        "開發工具、Clone Repository、補回 "
        ".env.local，再重新安裝 npm 套件。"
    ),
    "info",
)


# ------------------------------------------------------------
# 26.2
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.2 舊電腦移機前準備",
    level=2,
)

insert_numbered_before(
    version_heading,
    1,
    (
        "開啟 Visual Studio Code，確認所有"
        "程式修改都已儲存。"
    ),
)

insert_numbered_before(
    version_heading,
    2,
    (
        "進入 Source Control，確認重要修改"
        "都已 Commit。"
    ),
)

insert_numbered_before(
    version_heading,
    3,
    "執行 Sync Changes，將最新版本同步至 GitHub。",
)

insert_numbered_before(
    version_heading,
    4,
    "使用 Terminal 執行 git status。",
)

insert_code_before(
    version_heading,
    "git status",
)

insert_text_before(
    version_heading,
    "理想狀態應顯示：",
)

insert_code_before(
    version_heading,
    "nothing to commit, working tree clean",
)

insert_numbered_before(
    version_heading,
    5,
    (
        "備份 .env.local 至安全的私人位置，"
        "例如加密 USB、私人雲端或密碼管理器。"
    ),
)

insert_note_before(
    version_heading,
    "安全警告",
    (
        ".env.local 不可 Commit 到 GitHub。"
        "若舊電腦曾存放 service_role key、"
        "Token 或其他 Secret，也不可放入公開"
        " Repository。"
    ),
    "danger",
)


# ------------------------------------------------------------
# 26.3
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.3 不需要從舊電腦搬移的資料",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "下列資料可在新電腦重新產生，"
        "不需要從舊電腦直接複製："
    ),
)

for item in [
    "node_modules",
    ".next",
    "out",
    "npm Cache",
    "VS Code 暫存資料",
]:
    insert_bullet_before(
        version_heading,
        item,
    )

insert_note_before(
    version_heading,
    "建議",
    (
        "不要直接把整個 node_modules 從舊電腦"
        "複製到新電腦。應依 package.json 與 "
        "package-lock.json 在新電腦重新安裝套件。"
    ),
    "warning",
)


# ------------------------------------------------------------
# 26.4
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.4 新電腦安裝必要軟體",
    level=2,
)

insert_text_before(
    version_heading,
    "依序安裝：",
)

for item in [
    "Visual Studio Code",
    "Node.js 24",
    "Git",
    "Google Chrome 或 Microsoft Edge",
]:
    insert_bullet_before(
        version_heading,
        item,
    )

insert_text_before(
    version_heading,
    "安裝完成後開啟 PowerShell：",
)

insert_code_before(
    version_heading,
    """node --version
npm --version
git --version""",
)

insert_text_before(
    version_heading,
    (
        "Node.js 建議與 GitHub Actions 使用的"
        "主要版本一致，目前為 Node.js 24。"
    ),
)


# ------------------------------------------------------------
# 26.5
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.5 從 GitHub 還原主程式",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "若新電腦希望維持與原本相同的"
        "專案路徑，可直接執行："
    ),
)

insert_code_before(
    version_heading,
    (
        "git clone "
        "https://github.com/ShengHung-ken/"
        "titanium-it.git "
        r"D:\linni-computer"
    ),
)

insert_text_before(
    version_heading,
    "進入專案資料夾：",
)

insert_code_before(
    version_heading,
    r"cd D:\linni-computer",
)


# ------------------------------------------------------------
# 26.6
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.6 還原 .env.local",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "在專案根目錄重新建立："
    ),
)

insert_code_before(
    version_heading,
    r"D:\linni-computer\.env.local",
)

insert_text_before(
    version_heading,
    "內容格式：",
)

insert_code_before(
    version_heading,
    """NEXT_PUBLIC_SUPABASE_URL=你的 Supabase Project URL
NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY=你的 Supabase Publishable Key""",
)

insert_note_before(
    version_heading,
    "安全警告",
    (
        "這裡只放 Publishable Key。"
        "不得把 Supabase service_role key "
        "放進 NEXT_PUBLIC_* 或前端程式。"
    ),
    "danger",
)


# ------------------------------------------------------------
# 26.7
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.7 重新安裝 npm 套件",
    level=2,
)

insert_code_before(
    version_heading,
    "npm install",
)

insert_text_before(
    version_heading,
    (
        "npm 會依照 package.json 與 "
        "package-lock.json 自動還原專案所需套件。"
    ),
)


# ------------------------------------------------------------
# 26.8
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.8 本機功能測試",
    level=2,
)

insert_code_before(
    version_heading,
    "npm run dev",
)

insert_text_before(
    version_heading,
    "瀏覽器開啟：",
)

insert_code_before(
    version_heading,
    "http://localhost:3000/",
)

insert_text_before(
    version_heading,
    "依序確認：",
)

for item in [
    "首頁可以正常開啟。",
    "CSS 樣式正常。",
    "Shield Logo 正常。",
    "首頁大型 Logo 正常。",
    "商品資料可以從 Supabase 讀取。",
    "商品圖片正常。",
    "LINE 按鈕與 QR Code 正常。",
    "Facebook 按鈕正常。",
    "/login/ 可以正常開啟。",
    "管理員可以正常登入。",
    "/admin/ 可以正常管理商品。",
    "登出功能正常。",
]:
    insert_bullet_before(
        version_heading,
        item,
    )


# ------------------------------------------------------------
# 26.9
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.9 正式 Build 驗證",
    level=2,
)

insert_code_before(
    version_heading,
    "npm run build",
)

insert_text_before(
    version_heading,
    "應確認 Build 完整成功，例如：",
)

insert_code_before(
    version_heading,
    """✓ Compiled successfully
✓ Generating static pages
✓ Exporting""",
)

insert_note_before(
    version_heading,
    "重要",
    (
        "新電腦第一次接手專案時，"
        "必須至少完成一次 npm run build，"
        "不要只確認 npm run dev 可以啟動。"
    ),
    "warning",
)


# ------------------------------------------------------------
# 26.10
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.10 GitHub 登入與 Remote 確認",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "新電腦第一次使用 VS Code 的 "
        "Commit / Sync Changes 時，"
        "可能會要求重新登入 GitHub。"
    ),
)

insert_text_before(
    version_heading,
    "可使用以下指令確認 Remote：",
)

insert_code_before(
    version_heading,
    "git remote -v",
)

insert_text_before(
    version_heading,
    "應指向：",
)

insert_code_before(
    version_heading,
    (
        "https://github.com/ShengHung-ken/"
        "titanium-it.git"
    ),
)


# ------------------------------------------------------------
# 26.11
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.11 不需要重新設定的雲端服務",
    level=2,
)

insert_text_before(
    version_heading,
    (
        "更換開發電腦不代表網站重新建置。"
        "以下服務都位於雲端，一般不需要重新設定："
    ),
)

for item in [
    "GitHub Repository",
    "GitHub Actions",
    "GitHub Pages",
    "GitHub Repository Variables",
    "GitHub Pages Custom Domain",
    "titaniumit.rweb.site",
    "HTTPS Certificate",
    "Supabase Database",
    "Supabase Auth",
    "Supabase Storage",
    "Supabase RLS",
    "Google Search Console",
    "Google Sitemap",
    "Facebook 粉絲專頁",
    "LINE 官方帳號",
]:
    insert_bullet_before(
        version_heading,
        item,
    )


# ------------------------------------------------------------
# 26.12
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.12 新電腦還原完成檢查表",
    level=2,
)

check_items = [
    "Visual Studio Code 已安裝",
    "Node.js 24 已安裝",
    "npm 可正常執行",
    "Git 已安裝",
    "Repository Clone 成功",
    ".env.local 已建立",
    "npm install 成功",
    "npm run dev 成功",
    "首頁正常",
    "商品資料正常",
    "商品圖片正常",
    "管理員登入正常",
    "管理後台正常",
    "npm run build 成功",
    "GitHub Remote 正確",
    "VS Code 可以 Commit",
    "VS Code 可以 Sync Changes",
]

for item in check_items:
    insert_text_before(
        version_heading,
        f"☐ {item}",
    )


# ------------------------------------------------------------
# 26.13
# ------------------------------------------------------------

insert_heading_before(
    version_heading,
    "26.13 災難復原重點",
    level=2,
)

insert_code_before(
    version_heading,
    """GitHub
= 主程式與版本歷史

Supabase
= 商品資料 / 帳號 / 圖片

.env.local
= 新電腦需重新建立的本機設定

新電腦
= Clone + .env.local + npm install""",
)

insert_note_before(
    version_heading,
    "備份建議",
    (
        "除 GitHub 外，Word SOP 本身也應另外"
        "備份至私人 OneDrive、Google Drive、"
        "NAS 或 USB。若 SOP 沒有 Commit 到 GitHub，"
        "它只存在目前的本機硬碟中。"
    ),
    "info",
)


# ============================================================
# 將原本第 26 章改成第 27 章
# ============================================================

version_heading.text = (
    "第 27 章　版本修改紀錄"
)


# ============================================================
# 更新版本資訊
# ============================================================

update_cover_version(
    document
)

update_document_info_table(
    document
)

add_version_history(
    document
)

enable_field_update(
    document
)


# ============================================================
# 儲存 V1.1
# ============================================================

document.save(
    str(OUTPUT_FILE)
)


print()
print("SOP 更新完成：")
print(OUTPUT_FILE)
print()
print("新增章節：")
print(
    "第 26 章　新電腦移機 / 開發環境還原 SOP"
)
print()
print(
    "原本的「版本修改紀錄」已調整為第 27 章。"
)
print()
print(
    "開啟 Word 後請按 Ctrl+A，再按 F9，"
    "並選擇「更新整個目錄」。"
)