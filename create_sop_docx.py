from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = (
    BASE_DIR
    / "鈦鼎資訊_Titanium_IT_官網建置與維運SOP_V1.0.docx"
)
LOGO_FILE = BASE_DIR / "public" / "logo-shield.png"


document = Document()


# ============================================================
# 基本頁面設定
# ============================================================

for section in document.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)


# ============================================================
# 共用工具
# ============================================================

def set_run_font(
    run,
    chinese_font="Microsoft JhengHei",
    english_font="Aptos",
    size=None,
    bold=None,
    color=None,
):
    run.font.name = english_font

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        chinese_font,
    )

    if size is not None:
        run.font.size = Pt(size)

    if bold is not None:
        run.bold = bold

    if color is not None:
        run.font.color.rgb = RGBColor(
            *color,
        )


def set_paragraph_spacing(
    paragraph,
    before=0,
    after=6,
    line_spacing=1.25,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def style_normal():
    style = document.styles["Normal"]

    style.font.name = "Aptos"
    style.font.size = Pt(11)

    style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Microsoft JhengHei",
    )


def style_headings():
    heading_settings = {
        "Title": (28, True, (20, 45, 80)),
        "Heading 1": (18, True, (20, 60, 110)),
        "Heading 2": (15, True, (30, 80, 140)),
        "Heading 3": (13, True, (50, 90, 140)),
    }

    for style_name, (
        size,
        bold,
        color,
    ) in heading_settings.items():
        style = document.styles[style_name]

        style.font.name = "Aptos"

        style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Microsoft JhengHei",
        )

        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(
            *color,
        )


def add_heading(
    text,
    level=1,
):
    paragraph = document.add_heading(
        text,
        level=level,
    )

    set_paragraph_spacing(
        paragraph,
        before=8,
        after=8,
    )

    return paragraph


def add_text(
    text,
    bold=False,
    align=None,
):
    paragraph = document.add_paragraph()

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=11,
        bold=bold,
    )

    set_paragraph_spacing(
        paragraph,
        after=6,
    )

    return paragraph


def add_bullet(
    text,
    level=0,
):
    style_name = (
        "List Bullet"
        if level == 0
        else "List Bullet 2"
    )

    paragraph = document.add_paragraph(
        style=style_name,
    )

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=11,
    )

    set_paragraph_spacing(
        paragraph,
        after=3,
    )

    return paragraph


def add_numbered(
    text,
):
    paragraph = document.add_paragraph(
        style="List Number",
    )

    run = paragraph.add_run(text)

    set_run_font(
        run,
        size=11,
    )

    set_paragraph_spacing(
        paragraph,
        after=3,
    )

    return paragraph


def add_code_block(
    text,
):
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.right_indent = Cm(
        0.6
    )
    paragraph.paragraph_format.space_before = Pt(
        4
    )
    paragraph.paragraph_format.space_after = Pt(
        8
    )

    p_pr = paragraph._p.get_or_add_pPr()

    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        "F3F5F7",
    )
    p_pr.append(shading)

    border = OxmlElement("w:pBdr")

    for border_name in (
        "top",
        "left",
        "bottom",
        "right",
    ):
        item = OxmlElement(
            f"w:{border_name}"
        )
        item.set(
            qn("w:val"),
            "single",
        )
        item.set(
            qn("w:sz"),
            "4",
        )
        item.set(
            qn("w:color"),
            "D9E1E8",
        )
        border.append(item)

    p_pr.append(border)

    run = paragraph.add_run(text)

    set_run_font(
        run,
        chinese_font="Microsoft JhengHei",
        english_font="Consolas",
        size=9.5,
        color=(30, 40, 50),
    )

    return paragraph


def add_note(
    title,
    text,
    level="info",
):
    colors = {
        "info": ("EAF3FF", "2F6FB3"),
        "warning": ("FFF4CC", "B7791F"),
        "danger": ("FFE5E5", "B42318"),
    }

    fill_color, text_color = colors.get(
        level,
        colors["info"],
    )

    table = document.add_table(
        rows=1,
        cols=1,
    )

    table.autofit = True

    cell = table.cell(0, 0)
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    tc_pr = cell._tc.get_or_add_tcPr()

    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        fill_color,
    )
    tc_pr.append(shading)

    paragraph = cell.paragraphs[0]

    title_run = paragraph.add_run(
        f"{title}\n"
    )

    set_run_font(
        title_run,
        size=11,
        bold=True,
        color=tuple(
            int(
                text_color[index:index + 2],
                16,
            )
            for index in (0, 2, 4)
        ),
    )

    body_run = paragraph.add_run(text)

    set_run_font(
        body_run,
        size=10.5,
    )

    document.add_paragraph()


def add_table(
    headers,
    rows,
    widths=None,
):
    table = document.add_table(
        rows=1,
        cols=len(headers),
    )

    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for index, header in enumerate(headers):
        header_cells[index].text = header

        header_cells[
            index
        ].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        tc_pr = header_cells[
            index
        ]._tc.get_or_add_tcPr()

        shading = OxmlElement("w:shd")
        shading.set(
            qn("w:fill"),
            "D9EAF7",
        )
        tc_pr.append(shading)

        for paragraph in header_cells[
            index
        ].paragraphs:
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=10,
                    bold=True,
                )

    for row_data in rows:
        cells = table.add_row().cells

        for index, value in enumerate(
            row_data
        ):
            cells[index].text = str(value)

            cells[
                index
            ].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cells[
                index
            ].paragraphs:
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=10,
                    )

    if widths:
        for row in table.rows:
            for index, width in enumerate(
                widths
            ):
                row.cells[index].width = Cm(
                    width
                )

    document.add_paragraph()

    return table


def add_hyperlink(
    paragraph,
    text,
    url,
):
    part = paragraph.part

    relationship_id = (
        part.relate_to(
            url,
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/"
            "relationships/hyperlink",
            is_external=True,
        )
    )

    hyperlink = OxmlElement(
        "w:hyperlink"
    )

    hyperlink.set(
        qn("r:id"),
        relationship_id,
    )

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement(
        "w:rPr"
    )

    color = OxmlElement("w:color")
    color.set(
        qn("w:val"),
        "0563C1",
    )

    underline = OxmlElement("w:u")
    underline.set(
        qn("w:val"),
        "single",
    )

    run_properties.append(color)
    run_properties.append(underline)

    new_run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text

    new_run.append(text_element)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def add_link_line(
    label,
    url,
):
    paragraph = document.add_paragraph()

    run = paragraph.add_run(
        f"{label}："
    )

    set_run_font(
        run,
        size=11,
        bold=True,
    )

    add_hyperlink(
        paragraph,
        url,
        url,
    )

    set_paragraph_spacing(
        paragraph,
        after=5,
    )


def add_page_number(
    section,
):
    footer = section.footer

    paragraph = footer.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        "第 "
    )
    set_run_font(
        run,
        size=9,
    )

    fld_char_1 = OxmlElement(
        "w:fldChar"
    )
    fld_char_1.set(
        qn("w:fldCharType"),
        "begin",
    )

    instr_text = OxmlElement(
        "w:instrText"
    )
    instr_text.set(
        qn("xml:space"),
        "preserve",
    )
    instr_text.text = " PAGE "

    fld_char_2 = OxmlElement(
        "w:fldChar"
    )
    fld_char_2.set(
        qn("w:fldCharType"),
        "end",
    )

    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)

    end_run = paragraph.add_run(
        " 頁"
    )

    set_run_font(
        end_run,
        size=9,
    )


def add_toc():
    paragraph = document.add_paragraph()

    run = paragraph.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(
        qn("w:fldCharType"),
        "begin",
    )

    instr_text = OxmlElement(
        "w:instrText"
    )
    instr_text.set(
        qn("xml:space"),
        "preserve",
    )
    instr_text.text = (
        'TOC \\o "1-3" \\h \\z \\u'
    )

    fld_char_separate = OxmlElement(
        "w:fldChar"
    )
    fld_char_separate.set(
        qn("w:fldCharType"),
        "separate",
    )

    fld_char_end = OxmlElement(
        "w:fldChar"
    )
    fld_char_end.set(
        qn("w:fldCharType"),
        "end",
    )

    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(
        fld_char_separate
    )

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "請在 Word 中按 Ctrl+A，"
        "再按 F9 更新目錄。"
    )

    run._r.append(placeholder)
    run._r.append(fld_char_end)


def add_section_break():
    document.add_section(
        WD_SECTION.NEW_PAGE
    )


# ============================================================
# 初始化樣式
# ============================================================

style_normal()
style_headings()

for section in document.sections:
    add_page_number(section)


# ============================================================
# 封面
# ============================================================

document.add_paragraph()
document.add_paragraph()

if LOGO_FILE.exists():
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run()

    run.add_picture(
        str(LOGO_FILE),
        width=Inches(1.7),
    )


title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = title.add_run(
    "鈦鼎資訊 Titanium IT"
)

set_run_font(
    run,
    size=28,
    bold=True,
    color=(30, 55, 85),
)

subtitle = document.add_paragraph()
subtitle.alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)

run = subtitle.add_run(
    "官網建置、部署與維運標準作業程序"
)

set_run_font(
    run,
    size=20,
    bold=True,
    color=(40, 85, 130),
)

document.add_paragraph()

details = [
    "文件版本：V1.0",
    "建立日期：2026-08-28",
    "正式網站：https://titaniumit.rweb.site/",
    "專案名稱：Titanium IT 官方網站",
    (
        "文件用途：網站建置、部署、日常維護、"
        "故障排除及系統交接"
    ),
]

for item in details:
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(item)

    set_run_font(
        run,
        size=11,
    )


document.add_page_break()


# ============================================================
# 文件資訊
# ============================================================

add_heading(
    "文件資訊",
    level=1,
)

add_table(
    ["項目", "內容"],
    [
        [
            "文件名稱",
            "鈦鼎資訊 Titanium IT 官網建置與維運 SOP",
        ],
        ["文件版本", "V1.0"],
        ["建立日期", "2026-08-28"],
        [
            "正式網站",
            "https://titaniumit.rweb.site/",
        ],
        [
            "原 GitHub Pages",
            (
                "https://shenghung-ken.github.io/"
                "titanium-it/"
            ),
        ],
        [
            "GitHub Repository",
            "ShengHung-ken/titanium-it",
        ],
        ["前端框架", "Next.js 15"],
        [
            "程式語言",
            "TypeScript / React",
        ],
        ["CSS", "Tailwind CSS"],
        ["部署平台", "GitHub Pages"],
        [
            "Database",
            "Supabase PostgreSQL",
        ],
        [
            "Authentication",
            "Supabase Auth",
        ],
        [
            "Storage",
            "Supabase Storage",
        ],
        [
            "自訂網域",
            "titaniumit.rweb.site",
        ],
        [
            "SEO",
            "Google Search Console",
        ],
        [
            "開發作業系統",
            "Windows 10 / 11",
        ],
        [
            "編輯器",
            "Visual Studio Code",
        ],
        ["Runtime", "Node.js 24"],
        ["Package Manager", "npm"],
        [
            "Version Control",
            "Git / GitHub",
        ],
    ],
    widths=[5, 11],
)


add_heading(
    "目錄",
    level=1,
)

add_toc()

add_note(
    "目錄更新方式",
    (
        "第一次開啟 Word 後，請按 Ctrl+A 全選，"
        "再按 F9，即可更新自動目錄與頁碼。"
    ),
    "info",
)

document.add_page_break()


# ============================================================
# 第 1 章
# ============================================================

add_heading(
    "第 1 章　專案概述",
    level=1,
)

add_heading(
    "1.1 系統架構",
    level=2,
)

add_code_block(
    """使用者
   │
   ▼
https://titaniumit.rweb.site/
   │
   ▼
GitHub Pages
   │
   ▼
Next.js Static Export
   │
   ├──────────────► Supabase Database
   │
   ├──────────────► Supabase Auth
   │
   └──────────────► Supabase Storage"""
)

add_text(
    (
        "網站採 Next.js Static Export，"
        "GitHub Pages 負責提供靜態 HTML、CSS、"
        "JavaScript 與圖片；動態商品資料、登入與"
        "商品圖片則由 Supabase 提供。"
    )
)

add_heading(
    "1.2 正式網址",
    level=2,
)

add_link_line(
    "正式網站",
    "https://titaniumit.rweb.site/",
)

add_link_line(
    "原 GitHub Pages",
    (
        "https://shenghung-ken.github.io/"
        "titanium-it/"
    ),
)

add_text(
    (
        "原 GitHub Pages 網址目前已實測會自動"
        "導向新的自訂網域，因此過去分享出去的"
        "舊網址仍可繼續使用。"
    )
)

add_heading(
    "1.3 本機專案位置",
    level=2,
)

add_code_block(
    r"D:\linni-computer"
)

add_heading(
    "1.4 主要 Routes",
    level=2,
)

add_table(
    ["Route", "用途"],
    [
        ["/", "網站首頁"],
        ["/login/", "管理員登入"],
        ["/admin/", "商品管理後台"],
        [
            "/robots.txt",
            "搜尋引擎爬蟲規則",
        ],
        [
            "/sitemap.xml",
            "Google Sitemap",
        ],
    ],
)


# ============================================================
# 第 2 章
# ============================================================

add_heading(
    "第 2 章　建置環境與必要軟體",
    level=1,
)

add_heading(
    "2.1 必要軟體",
    level=2,
)

add_table(
    ["軟體", "用途", "建議版本"],
    [
        [
            "Windows",
            "開發作業系統",
            "Windows 10 / 11 64-bit",
        ],
        [
            "Visual Studio Code",
            "程式碼編輯",
            "最新 Stable",
        ],
        [
            "Node.js",
            "Next.js Runtime",
            "Node.js 24",
        ],
        [
            "npm",
            "套件管理",
            "隨 Node.js 安裝",
        ],
        [
            "Git",
            "版本控制",
            "最新 Stable",
        ],
        [
            "Chrome / Edge",
            "網站測試",
            "最新版本",
        ],
    ],
)

add_heading(
    "2.2 Visual Studio Code",
    level=2,
)

add_link_line(
    "下載",
    "https://code.visualstudio.com/",
)

add_text("建議安裝以下 Extension：")

for item in [
    "Tailwind CSS IntelliSense",
    "Prettier - Code formatter",
    "ESLint",
    "GitHub Pull Requests",
]:
    add_bullet(item)

add_note(
    "注意",
    (
        "不要同時安裝太多會自動格式化相同檔案的"
        "Extension，避免多個 Formatter 互相衝突。"
    ),
    "warning",
)

add_heading(
    "2.3 Node.js",
    level=2,
)

add_link_line(
    "官方網站",
    "https://nodejs.org/",
)

add_text(
    "目前 GitHub Actions 使用 Node.js 24。"
)

add_code_block(
    """node --version
npm --version"""
)

add_heading(
    "2.4 Git",
    level=2,
)

add_link_line(
    "官方網站",
    "https://git-scm.com/",
)

add_code_block(
    "git --version"
)

add_heading(
    "2.5 瀏覽器",
    level=2,
)

add_link_line(
    "Google Chrome",
    "https://www.google.com/chrome/",
)

add_link_line(
    "Microsoft Edge",
    "https://www.microsoft.com/edge",
)

add_text(
    (
        "正式部署後建議同時使用一般視窗、無痕視窗、"
        "桌面版及手機版進行驗收。"
    )
)

add_heading(
    "2.6 不需另外安裝的線上服務",
    level=2,
)

add_table(
    ["服務", "用途"],
    [
        [
            "GitHub",
            "Repository / Pages / Actions",
        ],
        [
            "Supabase",
            "Database / Auth / Storage",
        ],
        [
            "Google Search Console",
            "SEO / Sitemap / 索引",
        ],
        [
            "Schema Validator",
            "Structured Data 驗證",
        ],
        [
            "rweb.site",
            "免費自訂子網域",
        ],
        [
            "Cloudflare",
            "DNS 管理",
        ],
        [
            "EU.org",
            "免費網域申請",
        ],
        [
            "Facebook",
            "品牌粉絲專頁",
        ],
        [
            "LINE",
            "官方帳號聯絡",
        ],
        [
            "remove.bg",
            "圖片去背",
        ],
        [
            "Photopea",
            "圖片編輯",
        ],
    ],
)


# ============================================================
# 第 3 章
# ============================================================

add_heading(
    "第 3 章　專案取得與第一次安裝",
    level=1,
)

add_heading(
    "3.1 Clone Repository",
    level=2,
)

add_code_block(
    (
        "git clone "
        "https://github.com/ShengHung-ken/"
        "titanium-it.git\n"
        "cd titanium-it"
    )
)

add_heading(
    "3.2 安裝 npm 套件",
    level=2,
)

add_code_block(
    "npm install"
)

add_text("主要套件：")

for item in [
    "Next.js",
    "React",
    "TypeScript",
    "Tailwind CSS",
    "lucide-react",
    "@supabase/supabase-js",
]:
    add_bullet(item)

add_heading(
    "3.3 建立 .env.local",
    level=2,
)

add_code_block(
    """NEXT_PUBLIC_SUPABASE_URL=你的 Supabase Project URL
NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY=你的 Supabase Publishable Key"""
)

add_note(
    "安全警告",
    (
        ".env.local 不可 Commit 到 GitHub。"
        "service_role key、密碼及私人 Secret "
        "禁止放入 NEXT_PUBLIC_*。"
    ),
    "danger",
)


# ============================================================
# 第 4 章
# ============================================================

add_heading(
    "第 4 章　Next.js 專案設定",
    level=1,
)

add_heading(
    "4.1 next.config.mjs",
    level=2,
)

add_code_block(
    """/** @type {import('next').NextConfig} */
const nextConfig = {
  output: "export",

  trailingSlash: true,

  images: {
    unoptimized: true,
  },
};

export default nextConfig;"""
)

add_heading(
    "4.2 自訂網域路徑注意事項",
    level=2,
)

add_note(
    "重要",
    (
        "目前網站由 https://titaniumit.rweb.site/ "
        "根目錄提供，因此不要重新加入 "
        'basePath: "/titanium-it" 或 '
        'assetPrefix: "/titanium-it/"。'
    ),
    "warning",
)

add_text(
    "若重新加入上述設定，可能造成："
)

for item in [
    "CSS 404",
    "JavaScript 404",
    "Logo / 圖片 404",
    "網站只剩未套樣式的 HTML",
]:
    add_bullet(item)

add_heading(
    "4.3 Public 靜態資源",
    level=2,
)

add_table(
    ["檔案", "用途"],
    [
        [
            "public/logo-titanium.png",
            "首頁大型品牌圖與 SEO Logo",
        ],
        [
            "public/logo-shield.png",
            "頁首左上角透明盾牌 Logo",
        ],
        [
            (
                "public/"
                "google3dd4506f800c7c4d.html"
            ),
            "Google Search Console 驗證",
        ],
    ],
)

add_note(
    "重要",
    (
        "Google 驗證 HTML 檔案應保留，"
        "不要在網站驗證完成後刪除。"
    ),
    "warning",
)


# ============================================================
# 第 5 章
# ============================================================

add_heading(
    "第 5 章　本機開發與測試",
    level=1,
)

add_heading(
    "5.1 啟動開發環境",
    level=2,
)

add_code_block(
    "npm run dev"
)

add_link_line(
    "本機首頁",
    "http://localhost:3000/",
)

add_heading(
    "5.2 正式 Build",
    level=2,
)

add_code_block(
    "npm run build"
)

add_text(
    "成功時應看到類似："
)

add_code_block(
    """✓ Compiled successfully
✓ Linting and checking validity of types
✓ Generating static pages
✓ Exporting"""
)

add_text(
    "靜態輸出資料夾為："
)

add_code_block(
    "out"
)


# ============================================================
# 第 6 章
# ============================================================

add_heading(
    "第 6 章　GitHub 與自動部署",
    level=1,
)

add_link_line(
    "GitHub Repository",
    (
        "https://github.com/ShengHung-ken/"
        "titanium-it"
    ),
)

add_link_line(
    "GitHub Actions",
    (
        "https://github.com/ShengHung-ken/"
        "titanium-it/actions"
    ),
)

add_link_line(
    "GitHub Pages 設定",
    (
        "https://github.com/ShengHung-ken/"
        "titanium-it/settings/pages"
    ),
)

add_heading(
    "6.1 日常部署流程",
    level=2,
)

for item in [
    "儲存全部修改檔案。",
    "執行 npm run dev。",
    "完成本機功能與版面測試。",
    "執行 npm run build。",
    "確認 Build 完整成功。",
    "進入 VS Code Source Control。",
    "輸入 Commit Message。",
    "執行 Commit。",
    "執行 Sync Changes。",
    "等待 GitHub Actions 綠色成功。",
    "開正式網站進行部署後驗收。",
]:
    add_numbered(item)

add_heading(
    "6.2 GitHub Actions 重點",
    level=2,
)

add_code_block(
    """- name: Setup Pages
  uses: actions/configure-pages@v5"""
)

add_note(
    "重要",
    (
        "不要重新加入 static_site_generator: next。"
        "之前曾因 configure-pages 自動處理 Next.js "
        "而覆寫 basePath，導致部署路徑錯誤。"
    ),
    "warning",
)

add_heading(
    "6.3 Repository Variables",
    level=2,
)

add_code_block(
    """NEXT_PUBLIC_SUPABASE_URL
NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY"""
)


# ============================================================
# 第 7 章
# ============================================================

add_heading(
    "第 7 章　Supabase 後端系統",
    level=1,
)

add_link_line(
    "Supabase Dashboard",
    "https://supabase.com/dashboard",
)

add_text(
    "Supabase 在本專案負責："
)

for item in [
    "PostgreSQL Database",
    "商品資料",
    "Supabase Auth",
    "管理員登入",
    "Supabase Storage",
    "商品圖片",
]:
    add_bullet(item)

add_heading(
    "7.1 products 資料表",
    level=2,
)

add_table(
    ["欄位", "用途"],
    [
        ["id", "商品 ID"],
        ["name", "商品名稱"],
        ["category", "商品分類"],
        ["price", "售價"],
        ["stock", "庫存"],
        ["status", "上架 / 下架"],
        [
            "description",
            "商品描述陣列",
        ],
        [
            "image_url",
            "商品圖片 URL",
        ],
        [
            "created_at",
            "建立時間",
        ],
        [
            "updated_at",
            "更新時間",
        ],
    ],
)

add_heading(
    "7.2 前台商品規則",
    level=2,
)

add_text(
    (
        "一般訪客只允許讀取 status = 上架 的商品；"
        "管理員則可以讀取全部商品。"
    )
)


# ============================================================
# 第 8 章
# ============================================================

add_heading(
    "第 8 章　管理員登入與安全機制",
    level=1,
)

add_link_line(
    "登入頁",
    "https://titaniumit.rweb.site/login/",
)

add_link_line(
    "商品後台",
    "https://titaniumit.rweb.site/admin/",
)

add_heading(
    "8.1 管理員權限",
    level=2,
)

add_code_block(
    'app_metadata.role === "admin"'
)

add_text(
    (
        "前端登入後會檢查 Supabase User 的 "
        "app_metadata.role。"
    )
)

add_note(
    "安全警告",
    (
        "不要使用硬編碼 Email 作為真正的管理員"
        "授權條件。Email 可以作為登入資料，但真正"
        "的授權應使用 role 與 Supabase RLS。"
    ),
    "danger",
)

add_heading(
    "8.2 Row Level Security",
    level=2,
)

add_code_block(
    """public.is_admin()

auth.jwt()
→ app_metadata
→ role
→ admin"""
)

add_text(
    "目前權限原則："
)

for item in [
    "Public：只能 SELECT 上架商品。",
    "Admin：可 SELECT 全部商品。",
    "Admin：可 INSERT 商品。",
    "Admin：可 UPDATE 商品。",
    "Admin：可 DELETE 商品。",
]:
    add_bullet(item)

add_heading(
    "8.3 登入安全測試",
    level=2,
)

for item in [
    (
        "未登入直接開 /admin/，"
        "應自動導向 /login/。"
    ),
    "管理員登入後可進入 /admin/。",
    "登出後重新進 /admin/ 應再次要求登入。",
]:
    add_numbered(item)


# ============================================================
# 第 9 章
# ============================================================

add_heading(
    "第 9 章　商品管理系統",
    level=1,
)

add_heading(
    "9.1 新增商品",
    level=2,
)

add_text("後台新增商品需填入：")

for item in [
    "商品名稱",
    "商品分類",
    "售價",
    "庫存",
    "上下架狀態",
    "商品說明",
    "商品圖片",
]:
    add_bullet(item)

add_heading(
    "9.2 修改商品",
    level=2,
)

add_text(
    (
        "商品列表中的藍色鉛筆按鈕可進入編輯流程。"
        "修改完成後儲存即可。"
    )
)

add_heading(
    "9.3 刪除商品",
    level=2,
)

add_text(
    (
        "商品列表中的紅色垃圾桶按鈕用於刪除商品。"
        "刪除資料後，相關商品圖片亦應由 Storage "
        "清除。"
    )
)

add_heading(
    "9.4 日常商品維護",
    level=2,
)

add_note(
    "資訊",
    (
        "單純新增、修改、上下架或刪除商品不需要"
        "重新 Deploy 網站，因為前台會直接讀取"
        " Supabase 最新資料。"
    ),
    "info",
)


# ============================================================
# 第 10 章
# ============================================================

add_heading(
    "第 10 章　商品圖片 Storage",
    level=1,
)

add_text(
    "Supabase Storage Bucket："
)

add_code_block(
    "product-images"
)

add_text("權限：")

for item in [
    "Public Read",
    "Admin Upload",
    "Admin Update",
    "Admin Delete",
]:
    add_bullet(item)

add_heading(
    "10.1 圖片壓縮",
    level=2,
)

add_text(
    (
        "圖片上傳前會在瀏覽器端壓縮為 WebP，"
        "最大約 1600 × 1600，品質約 0.8。"
    )
)

add_code_block(
    """lib/image.ts
lib/product-storage.ts"""
)

add_heading(
    "10.2 圖片更新流程",
    level=2,
)

for item in [
    "先上傳新圖片。",
    "更新 Database。",
    (
        "Database 更新成功後，"
        "再刪除舊圖片。"
    ),
    (
        "若 Database 更新失敗，"
        "則刪除剛上傳的新圖片。"
    ),
]:
    add_numbered(item)


# ============================================================
# 第 11 章
# ============================================================

add_heading(
    "第 11 章　rweb.site 自訂網域",
    level=1,
)

add_link_line(
    "rweb.site",
    "https://rweb.site/",
)

add_link_line(
    "rweb.site GitHub",
    "https://github.com/katorlys/rweb.site",
)

add_text(
    "目前正式子網域："
)

add_code_block(
    "titaniumit.rweb.site"
)

add_heading(
    "11.1 申請流程",
    level=2,
)

for item in [
    "確認 titaniumit 未被占用。",
    "Fork katorlys/rweb.site。",
    "修改 records.json。",
    "加入 titaniumit 的 CNAME。",
    "Commit 變更。",
    "Open Pull Request。",
    "加入網站網址。",
    "加入網站首頁 Screenshot。",
    "完成 Checklist。",
    "等待 Validation / CodeRabbit。",
    "等待 Maintainer Review。",
    "PR Merged 後等待 DNS 生效。",
]:
    add_numbered(item)

add_heading(
    "11.2 CNAME 設定",
    level=2,
)

add_code_block(
    '"titaniumit": "shenghung-ken.github.io",'
)


# ============================================================
# 第 12 章
# ============================================================

add_heading(
    "第 12 章　GitHub Pages Custom Domain 與 HTTPS",
    level=1,
)

add_text(
    (
        "GitHub Repository → Settings → Pages "
        "→ Custom domain。"
    )
)

add_code_block(
    "titaniumit.rweb.site"
)

add_text("等待：")

add_code_block(
    "DNS Check in Progress"
)

add_text("成功後：")

add_code_block(
    "✓ DNS check successful"
)

add_heading(
    "12.1 HTTPS",
    level=2,
)

add_text(
    (
        "DNS 成功後，GitHub 需要時間簽發 "
        "TLS/SSL Certificate。"
    )
)

add_text("憑證完成後啟用：")

add_code_block(
    "☑ Enforce HTTPS"
)

add_heading(
    "12.2 舊網址導向",
    level=2,
)

add_text(
    (
        "目前原 GitHub Pages 網址已實測會"
        "自動導向 https://titaniumit.rweb.site/。"
    )
)


# ============================================================
# 第 13 章
# ============================================================

add_heading(
    "第 13 章　SEO 基礎設定",
    level=1,
)

add_text(
    (
        "主要 SEO Metadata 與 Structured Data "
        "位於 app/layout.tsx。"
    )
)

add_text("目前包含：")

for item in [
    "Title",
    "Description",
    "Keywords",
    "Canonical URL",
    "Robots",
    "Open Graph",
    "Twitter Card",
    "Icons",
    "Organization JSON-LD",
    "WebSite JSON-LD",
]:
    add_bullet(item)

add_heading(
    "13.1 正式 SEO Domain",
    level=2,
)

add_code_block(
    "https://titaniumit.rweb.site/"
)

add_heading(
    "13.2 Structured Data",
    level=2,
)

add_text(
    "目前 Structured Data 包含："
)

for item in [
    "Organization",
    "WebSite",
]:
    add_bullet(item)

add_link_line(
    "Schema Markup Validator",
    "https://validator.schema.org/",
)

add_heading(
    "13.3 robots.txt",
    level=2,
)

add_link_line(
    "robots.txt",
    (
        "https://titaniumit.rweb.site/"
        "robots.txt"
    ),
)

add_text(
    "目前排除："
)

add_code_block(
    """/admin/
/login/"""
)

add_heading(
    "13.4 sitemap.xml",
    level=2,
)

add_link_line(
    "sitemap.xml",
    (
        "https://titaniumit.rweb.site/"
        "sitemap.xml"
    ),
)


# ============================================================
# 第 14 章
# ============================================================

add_heading(
    "第 14 章　Google Search Console",
    level=1,
)

add_link_line(
    "Google Search Console",
    (
        "https://search.google.com/"
        "search-console/"
    ),
)

add_heading(
    "14.1 新增網站資源",
    level=2,
)

for item in [
    "進入 Google Search Console。",
    "選擇新增網站。",
    "選擇「網址前置字元」。",
    (
        "輸入 "
        "https://titaniumit.rweb.site/"
    ),
    "完成 HTML 檔案驗證。",
]:
    add_numbered(item)

add_heading(
    "14.2 驗證檔",
    level=2,
)

add_code_block(
    (
        "public/"
        "google3dd4506f800c7c4d.html"
    )
)

add_note(
    "重要",
    "此驗證檔不要刪除。",
    "warning",
)

add_heading(
    "14.3 Sitemap",
    level=2,
)

add_text(
    (
        "Search Console → Sitemap → "
        "輸入 sitemap.xml → 提交。"
    )
)

add_code_block(
    (
        "https://titaniumit.rweb.site/"
        "sitemap.xml"
    )
)

add_text(
    "目前狀態：成功，探索到 1 個公開網頁。"
)

add_heading(
    "14.4 要求建立索引",
    level=2,
)

for item in [
    "進入「網址審查」。",
    (
        "輸入 "
        "https://titaniumit.rweb.site/"
    ),
    "等待 Google 檢查。",
    "按「要求建立索引」。",
    "不要連續重複提交。",
]:
    add_numbered(item)

add_text(
    "可在 Google 搜尋："
)

add_code_block(
    "site:titaniumit.rweb.site"
)


# ============================================================
# 第 15 章
# ============================================================

add_heading(
    "第 15 章　LINE、Facebook 與聯絡資訊",
    level=1,
)

add_heading(
    "15.1 LINE",
    level=2,
)

add_link_line(
    "LINE 官方帳號",
    "https://lin.ee/PC2w13i",
)

add_text("網站目前包含：")

for item in [
    "LINE QR Code",
    "加入 LINE 官方帳號",
    "LINE 詢問",
]:
    add_bullet(item)

add_heading(
    "15.2 Facebook",
    level=2,
)

add_text(
    "Facebook Username："
)

add_code_block(
    "titaniumit.tw"
)

add_link_line(
    "Facebook 粉絲專頁",
    (
        "https://www.facebook.com/"
        "titaniumit.tw"
    ),
)

add_note(
    "注意",
    (
        "網站及 SEO 應使用正式 Facebook Username "
        "網址，不要使用 facebook.com/share/... "
        "分享網址。"
    ),
    "warning",
)

add_heading(
    "15.3 SEO sameAs",
    level=2,
)

add_code_block(
    """sameAs
├─ https://www.facebook.com/titaniumit.tw
└─ https://lin.ee/PC2w13i"""
)


# ============================================================
# 第 16 章
# ============================================================

add_heading(
    "第 16 章　公司名稱與商標查詢",
    level=1,
)

add_heading(
    "16.1 公司名稱",
    level=2,
)

add_link_line(
    "經濟部公司名稱預查輔助查詢",
    (
        "https://serv.gcis.nat.gov.tw/"
        "pub/cmpy/nameSearchListAction.do"
    ),
)

add_text(
    (
        "本次曾查詢「鈦鼎資訊」及「鈦鼎」。"
        "目前未看到完全同名的鈦鼎資訊，"
        "但正式能否使用仍應以正式名稱預查為準。"
    )
)

add_heading(
    "16.2 商標",
    level=2,
)

add_link_line(
    "智慧財產局商標檢索",
    (
        "https://cloud.tipo.gov.tw/"
        "S282/S282WV1/"
    ),
)

add_text("曾查詢：")

for item in [
    "鈦鼎資訊",
    "鈦鼎",
    "Titanium IT",
    "TITANIUM",
]:
    add_bullet(item)

add_text("相關分類：")

add_code_block(
    """009
035
037 / 3703
042"""
)

add_text(
    (
        "目前中文「鈦鼎資訊」風險相對較低；"
        "英文 TITANIUM IT 因存在較多 TITANIUM / "
        "TITAN 類近似商標，建議中文名稱作為"
        "主要品牌識別。"
    )
)


# ============================================================
# 第 17 章
# ============================================================

add_heading(
    "第 17 章　EU.org 與 Cloudflare 備案",
    level=1,
)

add_heading(
    "17.1 EU.org",
    level=2,
)

add_link_line(
    "EU.org",
    "https://nic.eu.org/",
)

add_text(
    "目前另申請："
)

add_code_block(
    "titaniumit.eu.org"
)

add_text(
    "目前狀態：等待 EU.org 人工審核。"
)

add_heading(
    "17.2 Cloudflare",
    level=2,
)

add_link_line(
    "Cloudflare Dashboard",
    "https://dash.cloudflare.com/",
)

add_text("目前 Nameserver：")

add_code_block(
    """sid.ns.cloudflare.com
zita.ns.cloudflare.com"""
)

add_note(
    "注意",
    (
        "EU.org 尚未核准以前，不要重複申請，"
        "也不要任意修改 Nameserver。"
    ),
    "warning",
)


# ============================================================
# 第 18 章
# ============================================================

add_heading(
    "第 18 章　日常維護 SOP",
    level=1,
)

for item in [
    "開啟 Visual Studio Code。",
    (
        r"開啟 D:\linni-computer。"
    ),
    "執行 npm run dev。",
    "本機確認功能正常。",
    "完成程式修改。",
    "儲存全部檔案。",
    "執行 npm run build。",
    "確認 Build 成功。",
    "Git Commit。",
    "Sync Changes。",
    "等待 GitHub Actions 成功。",
    "正式網站驗收。",
]:
    add_numbered(item)


# ============================================================
# 第 19 章
# ============================================================

add_heading(
    "第 19 章　網域切換 SOP",
    level=1,
)

add_text(
    (
        "若未來 titaniumit.eu.org 核准並決定"
        "取代 titaniumit.rweb.site，需修改以下項目。"
    )
)

for item in [
    "GitHub Pages Custom Domain",
    "app/layout.tsx",
    "app/robots.ts",
    "app/sitemap.ts",
]:
    add_bullet(item)

add_text(
    "修改後在 VS Code 全域搜尋："
)

add_code_block(
    "titaniumit.rweb.site"
)

add_text(
    "確認舊網址已全部更新後："
)

add_code_block(
    """npm run build
→ Deploy
→ DNS
→ HTTPS
→ Search Console
→ Sitemap
→ 要求建立索引"""
)


# ============================================================
# 第 20 章
# ============================================================

add_heading(
    "第 20 章　故障排除",
    level=1,
)

add_heading(
    "20.1 網站只有純文字、沒有 CSS",
    level=2,
)

add_text("優先檢查：")

add_code_block(
    """basePath
assetPrefix"""
)

add_text(
    "目前不應存在："
)

add_code_block(
    "/titanium-it"
)

add_heading(
    "20.2 Logo 404",
    level=2,
)

add_code_block(
    r"""Test-Path .\public\logo-titanium.png
Test-Path .\public\logo-shield.png
Test-Path .\out\logo-titanium.png
Test-Path .\out\logo-shield.png"""
)

add_heading(
    "20.3 DNS Check 一直 Pending",
    level=2,
)

for item in [
    "確認 rweb.site PR 是否已 Merge。",
    "確認 records.json CNAME 是否正確。",
    (
        "確認 GitHub Pages Custom Domain "
        "輸入正確。"
    ),
    (
        "不要因短暫 Pending 就反覆 "
        "Remove / Save。"
    ),
]:
    add_bullet(item)

add_heading(
    "20.4 HTTPS 無法勾選",
    level=2,
)

add_text(
    (
        "如果顯示 certificate has not yet been issued，"
        "代表 GitHub 尚在簽發 SSL Certificate，"
        "等待即可。"
    )
)

add_heading(
    "20.5 商品無法讀取",
    level=2,
)

for item in [
    "Supabase Project 狀態",
    "本機 .env.local",
    "GitHub Repository Variables",
    "Supabase RLS",
    "products table",
    "瀏覽器 Console",
]:
    add_bullet(item)


# ============================================================
# 第 21 章
# ============================================================

add_heading(
    "第 21 章　資安注意事項",
    level=1,
)

add_note(
    "高風險警告",
    (
        "Supabase service_role key 不可放入前端、"
        "GitHub Repository 或 NEXT_PUBLIC_*。"
    ),
    "danger",
)

add_text("絕對不要公開：")

for item in [
    "Supabase service_role key",
    "管理員密碼",
    "私人 Token",
    "Secret",
]:
    add_bullet(item)

add_text(
    (
        "NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY "
        "可用於前端；真正的資料安全必須由 "
        "Supabase Auth + RLS 提供。"
    )
)

add_note(
    "套件更新",
    (
        "不要未評估風險就執行 "
        "npm audit fix --force，"
        "避免自動升級造成 Breaking Changes。"
    ),
    "warning",
)


# ============================================================
# 第 22 章
# ============================================================

add_heading(
    "第 22 章　相關網站總表",
    level=1,
)

website_rows = [
    [
        "鈦鼎資訊正式網站",
        "https://titaniumit.rweb.site/",
    ],
    [
        "原 GitHub Pages",
        (
            "https://shenghung-ken.github.io/"
            "titanium-it/"
        ),
    ],
    [
        "GitHub Repository",
        (
            "https://github.com/ShengHung-ken/"
            "titanium-it"
        ),
    ],
    [
        "GitHub Actions",
        (
            "https://github.com/ShengHung-ken/"
            "titanium-it/actions"
        ),
    ],
    [
        "GitHub Pages Settings",
        (
            "https://github.com/ShengHung-ken/"
            "titanium-it/settings/pages"
        ),
    ],
    [
        "Supabase",
        "https://supabase.com/dashboard",
    ],
    [
        "Supabase Docs",
        "https://supabase.com/docs",
    ],
    [
        "rweb.site",
        "https://rweb.site/",
    ],
    [
        "rweb.site GitHub",
        "https://github.com/katorlys/rweb.site",
    ],
    [
        "Google Search Console",
        (
            "https://search.google.com/"
            "search-console/"
        ),
    ],
    [
        "Schema Validator",
        "https://validator.schema.org/",
    ],
    [
        "Facebook 粉絲專頁",
        (
            "https://www.facebook.com/"
            "titaniumit.tw"
        ),
    ],
    [
        "LINE 官方帳號",
        "https://lin.ee/PC2w13i",
    ],
    [
        "EU.org",
        "https://nic.eu.org/",
    ],
    [
        "Cloudflare",
        "https://dash.cloudflare.com/",
    ],
    [
        "公司名稱查詢",
        (
            "https://serv.gcis.nat.gov.tw/"
            "pub/cmpy/nameSearchListAction.do"
        ),
    ],
    [
        "商標查詢",
        (
            "https://cloud.tipo.gov.tw/"
            "S282/S282WV1/"
        ),
    ],
    [
        "Visual Studio Code",
        "https://code.visualstudio.com/",
    ],
    [
        "Node.js",
        "https://nodejs.org/",
    ],
    [
        "Git",
        "https://git-scm.com/",
    ],
    [
        "Google Chrome",
        "https://www.google.com/chrome/",
    ],
    [
        "remove.bg",
        "https://www.remove.bg/",
    ],
    [
        "Photopea",
        "https://www.photopea.com/",
    ],
]

table = document.add_table(
    rows=1,
    cols=2,
)

table.style = "Table Grid"

table.rows[0].cells[0].text = "用途"
table.rows[0].cells[1].text = "網站"

for cell in table.rows[0].cells:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(
        qn("w:fill"),
        "D9EAF7",
    )
    tc_pr.append(shading)

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(
                run,
                size=10,
                bold=True,
            )

for label, url in website_rows:
    cells = table.add_row().cells

    cells[0].text = label

    paragraph = cells[1].paragraphs[0]
    add_hyperlink(
        paragraph,
        url,
        url,
    )

    for run in cells[0].paragraphs[0].runs:
        set_run_font(
            run,
            size=10,
        )

document.add_paragraph()


# ============================================================
# 第 23 章
# ============================================================

add_heading(
    "第 23 章　常用指令速查",
    level=1,
)

command_rows = [
    [
        "啟動本機開發",
        "npm run dev",
    ],
    [
        "正式 Build",
        "npm run build",
    ],
    [
        "確認 Node",
        "node --version",
    ],
    [
        "確認 npm",
        "npm --version",
    ],
    [
        "確認 Git",
        "git --version",
    ],
    [
        "確認大型 Logo",
        (
            r"Test-Path "
            r".\public\logo-titanium.png"
        ),
    ],
    [
        "確認盾牌 Logo",
        (
            r"Test-Path "
            r".\public\logo-shield.png"
        ),
    ],
    [
        "確認 Build Logo",
        (
            r"Test-Path "
            r".\out\logo-titanium.png"
        ),
    ],
]

add_table(
    ["用途", "指令"],
    command_rows,
)


# ============================================================
# 第 24 章
# ============================================================

add_heading(
    "第 24 章　部署後驗收清單",
    level=1,
)

checklist = [
    "HTTPS 正常",
    "首頁 CSS 正常",
    "Shield Logo 正常",
    "首頁大型 Logo 正常",
    "手機版選單正常",
    "Supabase 商品正常",
    "商品圖片正常",
    "LINE QR Code 正常",
    "LINE 按鈕正常",
    "Facebook 按鈕正常",
    "Email 正常",
    "/login/ 正常",
    "管理員可登入",
    "/admin/ 正常",
    "未登入不可直接進入 /admin/",
    "登出正常",
    "robots.txt 正常",
    "sitemap.xml 正常",
    "舊 GitHub Pages 網址會自動轉址",
    "Google Search Console 正常",
]

for item in checklist:
    add_text(
        f"☐ {item}"
    )


# ============================================================
# 第 25 章
# ============================================================

add_heading(
    "第 25 章　禁止事項",
    level=1,
)

add_note(
    "高風險警告",
    (
        "以下事項可能造成資安風險、網站中斷、"
        "部署失敗或 SEO 異常。"
    ),
    "danger",
)

for item in [
    "刪除 Google 驗證 HTML 檔。",
    "Commit .env.local。",
    "公開 Supabase service_role key。",
    "關閉 Supabase RLS。",
    (
        "重新加入 /titanium-it "
        "basePath / assetPrefix。"
    ),
    (
        "用硬編碼管理員 Email "
        "作為真正授權機制。"
    ),
    (
        "沒有必要時修改或移除 "
        "GitHub Pages Custom Domain。"
    ),
    (
        "DNS Pending 時反覆重新設定。"
    ),
    (
        "使用 Facebook /share/... "
        "作為正式粉專網址。"
    ),
    (
        "未評估 Breaking Changes 就執行 "
        "npm audit fix --force。"
    ),
]:
    add_bullet(item)


# ============================================================
# 第 26 章
# ============================================================

add_heading(
    "第 26 章　版本修改紀錄",
    level=1,
)

add_table(
    ["版本", "日期", "修改內容"],
    [
        [
            "V1.0",
            "2026-08-28",
            (
                "完成 Next.js、GitHub Pages、"
                "Supabase、商品後台、rweb.site、"
                "HTTPS、SEO、Google Search Console、"
                "LINE、Facebook、商標及網域備案 SOP。"
            ),
        ],
    ],
)


# ============================================================
# 文件最後資訊
# ============================================================

document.add_page_break()

paragraph = document.add_paragraph()
paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)

run = paragraph.add_run(
    "鈦鼎資訊 Titanium IT"
)

set_run_font(
    run,
    size=18,
    bold=True,
    color=(30, 60, 100),
)

paragraph = document.add_paragraph()
paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)

run = paragraph.add_run(
    "官網建置與維運 SOP V1.0"
)

set_run_font(
    run,
    size=12,
)

paragraph = document.add_paragraph()
paragraph.alignment = (
    WD_ALIGN_PARAGRAPH.CENTER
)

run = paragraph.add_run(
    "End of Document"
)

set_run_font(
    run,
    size=9,
    color=(120, 120, 120),
)


# ============================================================
# Word 欄位更新設定
# ============================================================

settings = document.settings.element

update_fields = OxmlElement(
    "w:updateFields"
)

update_fields.set(
    qn("w:val"),
    "true",
)

settings.append(update_fields)


# ============================================================
# 儲存
# ============================================================

document.save(
    str(OUTPUT_FILE)
)

print()
print("Word SOP 已建立完成：")
print(OUTPUT_FILE)
print()
print(
    "開啟 Word 後，如目錄尚未更新，"
    "請按 Ctrl+A，再按 F9。"
)